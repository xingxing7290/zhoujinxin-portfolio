#!/usr/bin/env python3
"""Generate the three private, editable resume variants from public Markdown.

The Markdown sources intentionally contain no phone number or private email.
Both values must be supplied at runtime through RESUME_PHONE and RESUME_EMAIL;
generated files are written under the Git-ignored data/generated directory.
"""

from __future__ import annotations

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = ROOT / "data" / "generated" / "resumes"
VARIANTS = (
    ROOT / "resume" / "周金鑫-嵌入式软件工程师.md",
    ROOT / "resume" / "周金鑫-物联网全栈工程师.md",
    ROOT / "resume" / "周金鑫-AI原生全栈工程师.md",
)

# Design base: compact_reference_guide.
# Named overrides:
# - a4_resume_page: A4 portrait, 14 mm side and 12.5 mm vertical margins.
# - cjk_resume_type: Noto Sans SC and a compact 9 pt body for Chinese resumes.
# - customer_pack_resume_header: left-aligned name, target role, and one metadata line.
FONT = "Noto Sans SC"
NAVY = RGBColor(0x17, 0x2B, 0x4D)
BLUE = RGBColor(0x2F, 0x6B, 0x9A)
INK = RGBColor(0x20, 0x25, 0x2C)
MUTED = RGBColor(0x5A, 0x65, 0x72)
RULE = "D9E3EC"

PHONE_RE = re.compile(r"^1[3-9]\d{9}$")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def set_style_font(style, *, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    fonts.set(qn("w:cs"), FONT)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    fonts.set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color: str = RULE, size: str = "4") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def set_keep_with_next(paragraph, enabled: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = enabled
    paragraph.paragraph_format.keep_together = enabled


def ensure_style(doc: Document, name: str, base: str = "Normal"):
    styles = doc.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base]
    return style


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=9.0, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.8)
    normal.paragraph_format.line_spacing = 1.06
    normal.paragraph_format.widow_control = True

    title = ensure_style(doc, "Resume Title")
    set_style_font(title, size=23.5, color=NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.line_spacing = 1.0

    subtitle = ensure_style(doc, "Resume Subtitle")
    set_style_font(subtitle, size=10.8, color=BLUE, bold=True)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(1.5)
    subtitle.paragraph_format.line_spacing = 1.0

    contact = ensure_style(doc, "Resume Contact")
    set_style_font(contact, size=8.1, color=MUTED)
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(4.2)
    contact.paragraph_format.line_spacing = 1.0

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, size=11.4, color=NAVY, bold=True)
    h1.paragraph_format.space_before = Pt(4.6)
    h1.paragraph_format.space_after = Pt(2.2)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=9.6, color=INK, bold=True)
    h2.paragraph_format.space_before = Pt(2.8)
    h2.paragraph_format.space_after = Pt(0.8)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    bullet = ensure_style(doc, "Resume Bullet")
    set_style_font(bullet, size=8.75, color=INK)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(1.0)
    bullet.paragraph_format.line_spacing = 1.03
    bullet.paragraph_format.widow_control = True

    detail = ensure_style(doc, "Resume Detail")
    set_style_font(detail, size=8.45, color=MUTED)
    detail.paragraph_format.space_before = Pt(0.2)
    detail.paragraph_format.space_after = Pt(1.2)
    detail.paragraph_format.line_spacing = 1.02


def add_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(item.get(qn("w:abstractNumId")))
        for item in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "20")
    spacing.set(qn("w:line"), "247")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)

    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    rpr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    rpr.append(size)
    level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(num_id_el)


def add_inline_markdown(paragraph, text: str, *, base_color: RGBColor = INK) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, color=base_color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, color=base_color, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, color=base_color)


def add_field(paragraph, instruction: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7B8490")
    rpr.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "15")
    rpr.append(size)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def configure_page(doc: Document, name: str, role: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(12.5)
    section.bottom_margin = Mm(12.5)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(f"{name} · {role}    ")
    set_run_font(run, size=7.5, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" / ")
    set_run_font(run, size=7.5, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def extract_identity(lines: list[str], fallback_title: str) -> tuple[str, str]:
    for line in lines:
        if line.startswith("# "):
            parts = [item.strip() for item in line[2:].split("｜", 1)]
            name = parts[0]
            role = parts[1] if len(parts) > 1 else fallback_title
            return name, role
    return "周金鑫", fallback_title


def add_opening(doc: Document, name: str, role: str, phone: str, email: str, location: str) -> None:
    paragraph = doc.add_paragraph(style="Resume Title")
    set_keep_with_next(paragraph)
    add_inline_markdown(paragraph, name, base_color=NAVY)

    paragraph = doc.add_paragraph(style="Resume Subtitle")
    set_keep_with_next(paragraph)
    add_inline_markdown(paragraph, role, base_color=BLUE)

    contact = f"{location}  |  {phone}  |  {email}"
    paragraph = doc.add_paragraph(style="Resume Contact")
    set_keep_with_next(paragraph)
    add_inline_markdown(paragraph, contact, base_color=MUTED)


def add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 1")
    add_inline_markdown(paragraph, text, base_color=NAVY)
    set_paragraph_border_bottom(paragraph)


def add_subheading(doc: Document, text: str) -> None:
    parts = [part.strip() for part in text.split("｜")]
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Mm(181.5), WD_TAB_ALIGNMENT.RIGHT)
    title = parts[0]
    run = paragraph.add_run(title)
    set_run_font(run, size=9.6, color=INK, bold=True)

    if len(parts) == 2:
        run = paragraph.add_run(f"\t{parts[1]}")
        set_run_font(run, size=8.25, color=MUTED)
    elif len(parts) >= 3:
        middle = " · ".join(parts[1:-1])
        if middle:
            run = paragraph.add_run(f"  {middle}")
            set_run_font(run, size=8.25, color=MUTED)
        run = paragraph.add_run(f"\t{parts[-1]}")
        set_run_font(run, size=8.25, color=MUTED)


def build_resume(source: Path, output_dir: Path, phone: str, email: str, location: str) -> Path:
    lines = source.read_text(encoding="utf-8").splitlines()
    name, role = extract_identity(lines, source.stem.split("-", 1)[-1])

    doc = Document()
    configure_styles(doc)
    bullet_num_id = add_bullet_numbering(doc)
    configure_page(doc, name, role)
    add_opening(doc, name, role, phone, email, location)

    current_section = ""
    for raw in lines:
        line = raw.strip()
        if not line or line.startswith("# "):
            continue
        if line == "<!-- page-break -->":
            paragraph = doc.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            continue
        if line.startswith("## "):
            current_section = line[3:].strip()
            add_section_heading(doc, current_section)
            continue
        if line.startswith("### "):
            add_subheading(doc, line[4:].strip())
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="Resume Bullet")
            apply_bullet(paragraph, bullet_num_id)
            add_inline_markdown(paragraph, line[2:].strip())
            continue
        if line.startswith("> "):
            continue

        normalized_label = line.lstrip("*")
        detail_line = (
            normalized_label.startswith("阶段成果：")
            or normalized_label.startswith("技术栈：")
            or normalized_label.startswith("关键词：")
        )
        style = "Resume Detail" if detail_line else "Normal"
        paragraph = doc.add_paragraph(style=style)
        add_inline_markdown(paragraph, line, base_color=MUTED if detail_line else INK)

    doc.core_properties.title = f"{name}-{role}"
    doc.core_properties.subject = "个人简历"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / f"{source.stem}.docx"
    doc.save(output)
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate private DOCX resume variants.")
    parser.add_argument(
        "--source",
        action="append",
        type=Path,
        help="Markdown source to render. Repeat for multiple files; defaults to all variants.",
    )
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    phone = re.sub(r"\s+", "", os.getenv("RESUME_PHONE", ""))
    if not PHONE_RE.fullmatch(phone):
        raise SystemExit("RESUME_PHONE must be a valid private mainland China mobile number.")
    email = os.getenv("RESUME_EMAIL", "").strip()
    if not re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", email):
        raise SystemExit("RESUME_EMAIL must be supplied at runtime and contain a valid private email address.")
    location = os.getenv("RESUME_LOCATION", "北京").strip() or "北京"

    sources = tuple(args.source) if args.source else VARIANTS
    missing = [str(path) for path in sources if not path.is_file()]
    if missing:
        raise SystemExit("Missing resume source(s): " + ", ".join(missing))

    for source in sources:
        output = build_resume(source.resolve(), args.output_dir.resolve(), phone, email, location)
        print(output)


if __name__ == "__main__":
    main()
